from __future__ import annotations

import json
import io
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from icm_platform import db
from icm_platform.api import (
    _parse_requirement_document,
    app,
    load_cached_assertion_analysis,
    load_cached_report_analysis,
    load_report_analysis_versions,
    save_assertion_analysis,
    save_report_analysis,
)
from icm_platform.ai_service import AIConfigurationError, AIProviderError, AIService


class RequirementAITests(unittest.TestCase):
    def test_parse_requirement_docx_extracts_heading_and_table(self) -> None:
        from docx import Document

        document = Document()
        document.add_heading("用户登录需求", level=1)
        document.add_paragraph("用户输入账号密码后登录。")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "角色"
        table.rows[0].cells[1].text = "管理员"
        buffer = io.BytesIO()
        document.save(buffer)

        parsed = _parse_requirement_document("login.docx", buffer.getvalue())

        self.assertEqual(parsed["title"], "用户登录需求")
        self.assertIn("角色 | 管理员", parsed["text"])

    def test_spec_generation_payload_uses_requested_count_and_focus(self) -> None:
        payload = AIService()._spec_generation_payload(
            "MiniMax-M3",
            "登录需求",
            "标准",
            case_count=7,
            coverage_focus="abnormal_boundary",
        )
        user_content = json.loads(payload["messages"][1]["content"])

        self.assertIn("exactly 7", user_content["output_limits"])
        self.assertIn("异常输入", user_content["coverage_focus"])
        self.assertEqual(payload["max_completion_tokens"], 8192)

    def test_masks_saved_api_key(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with patch("icm_platform.db.DB_PATH", db_path), patch("icm_platform.db.DATA_DIR", db_path.parent):
                db.init_db()
                db.save_ai_settings(
                    {
                        "provider": "minimax-m3",
                        "base_url": "https://api.minimaxi.com/v1",
                        "model": "MiniMax-M3",
                        "api_key": "sk-cp-test-123456",
                    }
                )

                settings = db.get_ai_settings(mask_key=True)

        self.assertTrue(settings["has_api_key"])
        self.assertEqual(settings["api_key_masked"], "****3456")
        self.assertNotIn("api_key", settings)

    def test_normalizes_openai_compatible_base_url(self) -> None:
        self.assertEqual(
            AIService.chat_completions_url("https://api.minimaxi.com/v1"),
            "https://api.minimaxi.com/v1/chat/completions",
        )
        self.assertEqual(
            AIService.chat_completions_url("http://192.168.12.38:11434/v1"),
            "http://192.168.12.38:11434/v1/chat/completions",
        )

    def test_normalizes_ollama_tags_url(self) -> None:
        self.assertEqual(
            AIService.ollama_tags_url("http://192.168.12.38:11434/v1"),
            "http://192.168.12.38:11434/api/tags",
        )
        self.assertEqual(
            AIService.ollama_tags_url("http://192.168.12.38:11434/v1/chat/completions"),
            "http://192.168.12.38:11434/api/tags",
        )

    def test_parse_ollama_tags(self) -> None:
        raw = {
            "models": [
                {
                    "name": "qwen3.6:35b",
                    "model": "qwen3.6:35b",
                    "modified_at": "2026-06-01T07:14:42Z",
                    "size": 23938333577,
                    "details": {"parameter_size": "36.0B", "quantization_level": "Q4_K_M"},
                }
            ]
        }

        models = AIService.parse_ollama_tags(raw)

        self.assertEqual(models[0]["model"], "qwen3.6:35b")
        self.assertEqual(models[0]["details"]["parameter_size"], "36.0B")

    def test_parse_spec_cases_wraps_invalid_json(self) -> None:
        # PRD R6 / ARCH SK-01：错误文案统一为 "模型输出格式异常"，不保留旧英文串
        with self.assertRaisesRegex(AIProviderError, "模型输出格式异常"):
            AIService._parse_spec_cases('{"cases": [{"id": "CASE_001"} {"id": "CASE_002"}]}')

    def test_parse_spec_cases_accepts_json_object_fragment(self) -> None:
        parsed = AIService._parse_spec_cases('```json\n{"cases": [{"id": "CASE_001", "title": "ok"}]}\n```')

        self.assertEqual(parsed[0]["id"], "CASE_001")

    def test_parse_spec_cases_accepts_qwen_wrapped_json_with_schema_text(self) -> None:
        wrapped = (
            '下面先给出结果，字段格式参考 {"id":"MODULE_TYPE_NNN"}。\n'
            '```json\n'
            '{"cases":[{"id":"CASE_001","title":"ok"}]}\n'
            '```\n'
            '说明：以上 {schema} 仅供参考。'
        )

        parsed = AIService._parse_spec_cases(wrapped)

        self.assertEqual(parsed[0]["id"], "CASE_001")

    def test_spec_generation_payload_disables_ollama_thinking_and_raises_token_cap(self) -> None:
        service = AIService()

        payload = service._spec_generation_payload(
            "qwen3.6:35b",
            "需求：登录",
            "标准",
            provider="ollama-local",
        )

        self.assertEqual(payload["reasoning_effort"], "none")
        self.assertNotIn("think", payload)
        self.assertEqual(payload["max_tokens"], 12288)

    def test_generate_test_cases_spec_uses_longer_timeout_for_ollama(self) -> None:
        service = AIService()

        with patch.object(service, "_post_json", return_value={"choices": [{"message": {"content": '{"cases": []}'}}]}) as mocked:
            service.generate_test_cases_spec(
                "需求：登录",
                "标准",
                {
                    "provider": "ollama-local",
                    "api_key": "",
                    "base_url": "http://127.0.0.1:11434/v1",
                    "model": "qwen3.6:35b",
                },
            )

        self.assertEqual(mocked.call_args.kwargs["timeout"], 900)

    def test_generate_test_cases_spec_falls_back_to_reasoning_when_content_is_empty(self) -> None:
        service = AIService()
        raw = {
            "choices": [
                {
                    "message": {
                        "content": "",
                        "reasoning": '{"cases":[{"id":"CASE_001","title":"ok"}]}',
                    }
                }
            ]
        }

        with patch.object(service, "_post_json", return_value=raw):
            result = service.generate_test_cases_spec(
                "需求：登录",
                "标准",
                {
                    "provider": "ollama-local",
                    "api_key": "",
                    "base_url": "http://127.0.0.1:11434/v1",
                    "model": "qwen3.6:35b",
                },
            )

        self.assertEqual(result["cases"][0]["id"], "CASE_001")
        self.assertIn('"cases"', result["raw"])

    def test_generate_test_cases_spec_accepts_content_block_list(self) -> None:
        service = AIService()
        raw = {
            "choices": [
                {
                    "message": {
                        "content": [
                            {"type": "text", "text": '{"cases":[{"id":"CASE_001","title":"ok"}]}'},
                        ],
                    }
                }
            ]
        }

        with patch.object(service, "_post_json", return_value=raw):
            result = service.generate_test_cases_spec(
                "需求：登录",
                "标准",
                {
                    "provider": "ollama-local",
                    "api_key": "",
                    "base_url": "http://127.0.0.1:11434/v1",
                    "model": "qwen3.6:35b",
                },
            )

        self.assertEqual(result["cases"][0]["id"], "CASE_001")

    def test_generate_test_cases_spec_accepts_reasoning_object(self) -> None:
        service = AIService()
        raw = {
            "choices": [
                {
                    "message": {
                        "content": "",
                        "reasoning": {"text": '{"cases":[{"id":"CASE_001","title":"ok"}]}'},
                    }
                }
            ]
        }

        with patch.object(service, "_post_json", return_value=raw):
            result = service.generate_test_cases_spec(
                "需求：登录",
                "标准",
                {
                    "provider": "ollama-local",
                    "api_key": "",
                    "base_url": "http://127.0.0.1:11434/v1",
                    "model": "qwen3.6:35b",
                },
            )

        self.assertEqual(result["cases"][0]["id"], "CASE_001")

    def test_parse_chat_completion_json_response(self) -> None:
        content = {
            "test_points": [
                {
                    "name": "请求协助入口",
                    "category": "功能",
                    "priority": "P0",
                    "status": "待确认",
                    "description": "设备信息卡片中可发起请求协助",
                }
            ],
            "analysis_summary": "覆盖核心远程报修入口。",
            "risk_summary": "需要等待设备屏幕加载完成。",
            "case_count": 1,
        }
        raw = {"choices": [{"message": {"content": json.dumps(content, ensure_ascii=False)}}]}

        parsed = AIService.parse_chat_completion(raw)

        self.assertEqual(parsed["test_points"][0].name, "请求协助入口")
        self.assertEqual(parsed["analysis_summary"], "覆盖核心远程报修入口。")
        self.assertEqual(parsed["case_count"], 1)

    def test_minimax_requires_api_key(self) -> None:
        service = AIService()

        with self.assertRaises(AIConfigurationError):
            service._validate_settings(
                {
                    "provider": "minimax-m3",
                    "api_key": "",
                    "base_url": "https://api.minimaxi.com/v1",
                    "model": "MiniMax-M3",
                }
            )

    def test_ollama_provider_does_not_require_api_key(self) -> None:
        service = AIService()

        service._validate_settings(
            {
                "provider": "ollama-local",
                "api_key": "",
                "base_url": "http://192.168.12.38:11434/v1",
                "model": "qwen3.6:35b",
            }
        )

    def test_model_request_timeout_is_provider_error(self) -> None:
        service = AIService()

        with patch("urllib.request.urlopen", side_effect=TimeoutError("timed out")):
            with self.assertRaisesRegex(AIProviderError, "timed out after 300 seconds"):
                service._post_json("http://example.test/v1/chat/completions", "", {"model": "qwen3.6:35b"}, timeout=300)

    def test_omits_authorization_when_api_key_is_empty(self) -> None:
        self.assertEqual(AIService.request_headers(""), {"Content-Type": "application/json"})

    def test_ollama_provider_ignores_saved_api_key(self) -> None:
        self.assertEqual(
            AIService.api_key_for_provider({"provider": "ollama-local", "api_key": "sk-cp-old"}),
            "",
        )

    def test_analyze_run_report_with_ai_parses_json(self) -> None:
        service = AIService()
        raw = {
            "choices": [
                {
                    "message": {
                        "content": json.dumps(
                            {
                                "status": "failed",
                                "conclusion": "远程桌面未打开，需复查入口点击。",
                                "risks": ["截图缺少 final 阶段"],
                                "retest_suggestions": ["复跑 TC-ICM-012 并观察远程页签"],
                            },
                            ensure_ascii=False,
                        )
                    }
                }
            ]
        }

        with patch.object(service, "_post_json", return_value=raw):
            analysis = service.analyze_run_report_with_ai(
                "status: failed\nerror: remote desktop missing",
                [{"filename": "03-final.png", "case_id": "TC-ICM-012"}],
                ["open report", "missing tab"],
                {
                    "provider": "ollama-local",
                    "api_key": "",
                    "base_url": "http://127.0.0.1:11434/v1",
                    "model": "qwen3.6:35b",
                },
            )

        self.assertEqual(analysis["source"], "ai")
        self.assertEqual(analysis["status"], "failed")
        self.assertEqual(analysis["provider"], "ollama-local")
        self.assertEqual(analysis["model"], "qwen3.6:35b")
        self.assertEqual(analysis["risks"], ["截图缺少 final 阶段"])

    def test_report_analysis_accepts_ui_field_aliases(self) -> None:
        service = AIService()
        raw = {
            "choices": [
                {
                    "message": {
                        "content": json.dumps(
                            {
                                "status": "passed",
                                "current_analysis": "用例已按预期通过。",
                                "risk_tips": ["缺少最终页面截图"],
                                "retest_advice": ["补充一次带截图的复测"],
                            },
                            ensure_ascii=False,
                        )
                    }
                }
            ]
        }

        with patch.object(service, "_post_json", return_value=raw):
            analysis = service.analyze_run_report_with_ai(
                "status: passed",
                [],
                [],
                {
                    "provider": "ollama-local",
                    "api_key": "",
                    "base_url": "http://127.0.0.1:11434/v1",
                    "model": "qwen3.6:35b",
                },
            )

        self.assertEqual(analysis["conclusion"], "用例已按预期通过。")
        self.assertEqual(analysis["risks"], ["缺少最终页面截图"])
        self.assertEqual(analysis["retest_suggestions"], ["补充一次带截图的复测"])

    def test_switching_to_ollama_clears_saved_key(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with patch("icm_platform.db.DB_PATH", db_path), patch("icm_platform.db.DATA_DIR", db_path.parent):
                db.init_db()
                db.save_ai_settings(
                    {
                        "provider": "minimax-m3",
                        "base_url": "https://api.minimaxi.com/v1",
                        "model": "MiniMax-M3",
                        "api_key": "sk-cp-test-123456",
                    }
                )
                db.save_ai_settings(
                    {
                        "provider": "ollama-local",
                        "base_url": "http://192.168.12.38:11434/v1",
                        "model": "qwen3.6:35b",
                    }
                )

                settings = db.get_ai_settings(mask_key=True)

        self.assertFalse(settings["has_api_key"])
        self.assertEqual(settings["api_key_masked"], "")

    def test_report_analysis_cache_roundtrip(self) -> None:
        settings = {
            "provider": "ollama-local",
            "api_key": "",
            "base_url": "http://127.0.0.1:11434/v1",
            "model": "qwen3.6:35b",
        }
        analysis = {
            "provider": "ollama-local",
            "model": "qwen3.6:35b",
            "source": "ai",
            "status": "passed",
            "conclusion": "cached result",
            "risks": [],
            "retest_suggestions": [],
            "screenshot_count": 3,
            "log_count": 2,
        }
        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with patch("icm_platform.db.DB_PATH", db_path), patch("icm_platform.db.DATA_DIR", db_path.parent):
                db.init_db()
                save_report_analysis("run-1", "status: passed", settings, analysis)
                cached = load_cached_report_analysis("run-1", "status: passed", settings)
                missed = load_cached_report_analysis("run-1", "status: changed", settings)

        self.assertIsNotNone(cached)
        self.assertEqual(cached["conclusion"], "cached result")
        self.assertTrue(cached["cached"])
        self.assertIsNone(missed)

    def test_report_analysis_versions_keep_history(self) -> None:
        settings = {
            "provider": "ollama-local",
            "api_key": "",
            "base_url": "http://127.0.0.1:11434/v1",
            "model": "qwen3.6:35b",
        }
        first = {
            "provider": "ollama-local",
            "model": "qwen3.6:35b",
            "source": "ai",
            "status": "failed",
            "conclusion": "first result",
            "risks": [],
            "retest_suggestions": [],
            "screenshot_count": 1,
            "log_count": 1,
        }
        second = {**first, "conclusion": "second result"}
        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with patch("icm_platform.db.DB_PATH", db_path), patch("icm_platform.db.DATA_DIR", db_path.parent):
                db.init_db()
                save_report_analysis("run-2", "status: failed", settings, first)
                save_report_analysis("run-2", "status: failed", settings, second)
                cached = load_cached_report_analysis("run-2", "status: failed", settings)
                versions = load_report_analysis_versions("run-2", "status: failed")

        self.assertEqual(cached["conclusion"], "second result")
        self.assertEqual(len(versions), 2)
        self.assertEqual(versions[0]["analysis"]["conclusion"], "second result")
        self.assertEqual(versions[1]["analysis"]["conclusion"], "first result")

    def test_analyze_report_endpoint_runs_ai_and_saves_version(self) -> None:
        try:
            from fastapi.testclient import TestClient
        except Exception as exc:  # pragma: no cover
            self.skipTest(f"TestClient unavailable: {exc}")

        settings = {
            "provider": "ollama-local",
            "base_url": "http://127.0.0.1:11434/v1",
            "model": "qwen3.6:35b",
        }
        analysis = {
            "provider": "ollama-local",
            "model": "qwen3.6:35b",
            "source": "ai",
            "status": "failed",
            "conclusion": "AI 已完成报告解读",
            "risks": ["断言截图不足"],
            "retest_suggestions": ["重新执行失败步骤"],
            "screenshot_count": 1,
            "log_count": 0,
        }
        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with (
                patch("icm_platform.db.DB_PATH", db_path),
                patch("icm_platform.db.DATA_DIR", db_path.parent),
                patch("icm_platform.api.read_report", return_value="status: failed\nerror: assert failed"),
                patch("icm_platform.api.parse_report", return_value={"screenshots": [{"filename": "final.png"}], "case_id": "TC-ICM-001"}),
                patch("icm_platform.api.ai_service.analyze_run_report_with_ai", return_value=analysis) as analyze,
            ):
                db.init_db()
                db.save_ai_settings(settings)
                client = TestClient(app)
                response = client.post("/api/reports/run-9/analyze", json={"force": True})
                versions = client.get("/api/reports/run-9/analyses")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["conclusion"], "AI 已完成报告解读")
        self.assertFalse(response.json()["cached"])
        self.assertEqual(versions.status_code, 200)
        self.assertEqual(versions.json()[0]["analysis"]["conclusion"], "AI 已完成报告解读")
        analyze.assert_called_once()

    def test_analyze_report_endpoint_falls_back_to_run_detail_without_markdown(self) -> None:
        try:
            from fastapi.testclient import TestClient
        except Exception as exc:  # pragma: no cover
            self.skipTest(f"TestClient unavailable: {exc}")

        settings = {
            "provider": "ollama-local",
            "base_url": "http://127.0.0.1:11434/v1",
            "model": "qwen3.6:35b",
        }
        detail = {
            "run_id": "run-no-md",
            "case_id": "ICMDEV_BND_007",
            "case_name": "ICMDEV_BND_007 - 执行报告",
            "mode": "agent",
            "status": "completed",
            "started_at": "2026-06-23 10:00:00",
            "finished_at": "2026-06-23 10:01:00",
            "final_url": "https://example.test/#/hubble/device",
            "summary": {"conclusion": "执行完成", "failure_reason": ""},
            "steps": [{"step_index": 1, "title": "点击确定", "status": "completed", "summary": "弹窗关闭"}],
            "screenshots": [{"filename": "final.png", "case_id": "ICMDEV_BND_007"}],
        }
        analysis = {
            "provider": "ollama-local",
            "model": "qwen3.6:35b",
            "source": "ai",
            "status": "passed",
            "conclusion": "结构化报告已分析",
            "risks": [],
            "retest_suggestions": [],
            "screenshot_count": 1,
            "log_count": 0,
        }
        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with (
                patch("icm_platform.db.DB_PATH", db_path),
                patch("icm_platform.db.DATA_DIR", db_path.parent),
                patch("icm_platform.api.read_report", side_effect=FileNotFoundError("run-no-md")),
                patch("icm_platform.api._build_unified_run_detail", return_value=detail),
                patch("icm_platform.api.ai_service.analyze_run_report_with_ai", return_value=analysis),
            ):
                db.init_db()
                db.save_ai_settings(settings)
                client = TestClient(app)
                response = client.post("/api/reports/run-no-md/analyze", json={"force": True})
                versions = client.get("/api/reports/run-no-md/analyses")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["conclusion"], "结构化报告已分析")
        self.assertEqual(versions.status_code, 200)
        self.assertEqual(versions.json()[0]["analysis"]["conclusion"], "结构化报告已分析")


    # ===== E.1 AI 断言解析单元测试 =====

    def test_parse_assertions_with_ai_returns_structured_list(self) -> None:
        service = AIService()
        raw = {"choices": [{"message": {"content": json.dumps({"assertions": [
            {"type": "page_title_contains", "expected": "configure", "match_mode": "contains", "label": "页面标题"},
        ]})}}]}
        with patch.object(service, "_post_json", return_value=raw):
            result = service.parse_assertions_with_ai("页面标题包含 configure", {
                "provider": "minimax-m3", "base_url": "https://api.minimaxi.com/v1",
                "model": "MiniMax-M3", "api_key": "sk-test",
            })
        self.assertEqual(len(result), 1)
        item = result[0]
        self.assertEqual(item["type"], "page_title_contains")
        self.assertEqual(item["expected"], "configure")
        self.assertEqual(item["match_mode"], "contains")
        self.assertEqual(item["status"], "queued")
        self.assertEqual(item["source"], "ai")

    def test_parse_assertions_with_ai_normalizes_missing_match_mode(self) -> None:
        service = AIService()
        raw = {"choices": [{"message": {"content": json.dumps({"assertions": [
            {"type": "text_contains", "expected": "hello"},
        ]})}}]}
        with patch.object(service, "_post_json", return_value=raw):
            result = service.parse_assertions_with_ai("hello", {
                "provider": "minimax-m3", "base_url": "https://api.minimaxi.com/v1",
                "model": "MiniMax-M3", "api_key": "sk-test",
            })
        self.assertEqual(result[0]["match_mode"], "contains")

    def test_parse_assertions_with_ai_normalizes_unknown_type_to_text_contains(self) -> None:
        service = AIService()
        raw = {"choices": [{"message": {"content": json.dumps({"assertions": [
            {"type": "", "expected": "something"},
        ]})}}]}
        with patch.object(service, "_post_json", return_value=raw):
            result = service.parse_assertions_with_ai("something", {
                "provider": "minimax-m3", "base_url": "https://api.minimaxi.com/v1",
                "model": "MiniMax-M3", "api_key": "sk-test",
            })
        self.assertEqual(result[0]["type"], "text_contains")

    def test_parse_assertions_with_ai_raises_on_empty_assertions(self) -> None:
        service = AIService()
        raw = {"choices": [{"message": {"content": json.dumps({"assertions": []})}}]}
        with patch.object(service, "_post_json", return_value=raw):
            with self.assertRaises(AIProviderError):
                service.parse_assertions_with_ai("test", {
                    "provider": "minimax-m3", "base_url": "https://api.minimaxi.com/v1",
                    "model": "MiniMax-M3", "api_key": "sk-test",
                })

    def test_parse_assertions_with_ai_raises_on_invalid_json(self) -> None:
        service = AIService()
        raw = {"choices": [{"message": {"content": "not json at all"}}]}
        with patch.object(service, "_post_json", return_value=raw):
            with self.assertRaises(AIProviderError):
                service.parse_assertions_with_ai("test", {
                    "provider": "minimax-m3", "base_url": "https://api.minimaxi.com/v1",
                    "model": "MiniMax-M3", "api_key": "sk-test",
                })

    def test_parse_assertions_with_ai_raises_on_missing_content(self) -> None:
        service = AIService()
        raw = {}
        with patch.object(service, "_post_json", return_value=raw):
            with self.assertRaises(AIProviderError):
                service.parse_assertions_with_ai("test", {
                    "provider": "minimax-m3", "base_url": "https://api.minimaxi.com/v1",
                    "model": "MiniMax-M3", "api_key": "sk-test",
                })

    def test_parse_assertions_with_ai_validates_settings_first(self) -> None:
        service = AIService()
        with patch.object(service, "_post_json") as mocked:
            with self.assertRaises(AIConfigurationError):
                service.parse_assertions_with_ai("test", {
                    "provider": "minimax-m3", "base_url": "https://api.minimaxi.com/v1",
                    "model": "MiniMax-M3", "api_key": "",
                })
        mocked.assert_not_called()

    def test_assertion_parsing_payload_disables_ollama_thinking(self) -> None:
        service = AIService()
        payload = service._assertion_parsing_payload("qwen3.6:35b", "test", "ollama-local")
        self.assertEqual(payload.get("reasoning_effort"), "none")
        self.assertNotIn("thinking", payload)

    # ===== E.2 缓存命中/未命中测试 =====

    def test_assertion_analysis_cache_roundtrip(self) -> None:
        settings = {"provider": "minimax-m3", "model": "MiniMax-M3", "base_url": "", "api_key": ""}
        assertions = [{"type": "text_contains", "expected": "hello", "label": "文本", "status": "completed", "actual": "hello", "match_mode": "contains", "source": "ai"}]
        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with patch("icm_platform.db.DB_PATH", db_path), patch("icm_platform.db.DATA_DIR", db_path.parent):
                db.init_db()
                save_assertion_analysis("hello", settings, assertions)
                cached = load_cached_assertion_analysis("hello", settings)
        self.assertIsNotNone(cached)
        self.assertEqual(len(cached), 1)
        self.assertEqual(cached[0]["status"], "queued")
        self.assertEqual(cached[0]["actual"], "")
        self.assertTrue(cached[0]["cached"])

    def test_assertion_analysis_cache_miss_on_different_text(self) -> None:
        settings = {"provider": "minimax-m3", "model": "MiniMax-M3", "base_url": "", "api_key": ""}
        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with patch("icm_platform.db.DB_PATH", db_path), patch("icm_platform.db.DATA_DIR", db_path.parent):
                db.init_db()
                save_assertion_analysis("text A", settings, [{"type": "text_contains", "expected": "A"}])
                missed = load_cached_assertion_analysis("text B", settings)
        self.assertIsNone(missed)

    def test_assertion_analysis_cache_miss_on_different_model(self) -> None:
        settings_a = {"provider": "minimax-m3", "model": "MiniMax-M3", "base_url": "", "api_key": ""}
        settings_b = {"provider": "minimax-m3", "model": "MiniMax-M2", "base_url": "", "api_key": ""}
        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with patch("icm_platform.db.DB_PATH", db_path), patch("icm_platform.db.DATA_DIR", db_path.parent):
                db.init_db()
                save_assertion_analysis("same text", settings_a, [{"type": "text_contains", "expected": "x"}])
                missed = load_cached_assertion_analysis("same text", settings_b)
        self.assertIsNone(missed)

    def test_assertion_analysis_cache_corrupt_json_returns_none(self) -> None:
        settings = {"provider": "minimax-m3", "model": "MiniMax-M3", "base_url": "", "api_key": ""}
        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with patch("icm_platform.db.DB_PATH", db_path), patch("icm_platform.db.DATA_DIR", db_path.parent):
                db.init_db()
                with db.connect() as conn:
                    conn.execute(
                        "insert into assertion_analyses(expected_hash, provider, model, assertions_json, created_at, updated_at) values (?, ?, ?, ?, ?, ?)",
                        ("deadbeef", "minimax-m3", "MiniMax-M3", "NOT JSON", "2026-01-01T00:00:00Z", "2026-01-01T00:00:00Z"),
                    )
                result = load_cached_assertion_analysis("whatever", settings)
        self.assertIsNone(result)


    def test_reports_endpoint_hides_logically_deleted_runs(self) -> None:
        try:
            from fastapi.testclient import TestClient
        except Exception as exc:  # pragma: no cover
            self.skipTest(f"TestClient unavailable: {exc}")

        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with (
                patch("icm_platform.db.DB_PATH", db_path),
                patch("icm_platform.db.DATA_DIR", db_path.parent),
                patch(
                    "icm_platform.api.list_reports",
                    return_value=[
                        {
                            "run_id": "run-visible",
                            "case_id": "TC-001",
                            "case_name": "Visible case",
                            "status": "completed",
                            "path": "reports/run-visible.md",
                            "updated_at": 1,
                            "screenshot_count": 1,
                        },
                        {
                            "run_id": "run-deleted",
                            "case_id": "TC-002",
                            "case_name": "Deleted case",
                            "status": "failed",
                            "path": "reports/run-deleted.md",
                            "updated_at": 2,
                            "screenshot_count": 2,
                        },
                    ],
                ),
            ):
                db.init_db()
                now = "2026-07-08T10:00:00Z"
                with db.connect() as conn:
                    conn.execute(
                        "insert into run_tasks(id, mode, case_id, status, command, created_at, started_at, finished_at, report_deleted_at) values (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                        ("run-visible", "agent-explore", "TC-001", "completed", "run", now, now, now, None),
                    )
                    conn.execute(
                        "insert into run_tasks(id, mode, case_id, status, command, created_at, started_at, finished_at, report_deleted_at) values (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                        ("run-deleted", "agent-explore", "TC-002", "failed", "run", now, now, now, now),
                    )
                client = TestClient(app)
                response = client.get("/api/reports")

        self.assertEqual(response.status_code, 200)
        self.assertEqual([item["run_id"] for item in response.json()], ["run-visible"])

    def test_delete_report_endpoint_marks_report_as_deleted(self) -> None:
        try:
            from fastapi.testclient import TestClient
        except Exception as exc:  # pragma: no cover
            self.skipTest(f"TestClient unavailable: {exc}")

        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with patch("icm_platform.db.DB_PATH", db_path), patch("icm_platform.db.DATA_DIR", db_path.parent):
                db.init_db()
                now = "2026-07-08T10:00:00Z"
                with db.connect() as conn:
                    conn.execute(
                        "insert into run_tasks(id, mode, case_id, status, command, created_at, started_at, finished_at) values (?, ?, ?, ?, ?, ?, ?, ?)",
                        ("run-delete", "agent-explore", "TC-003", "completed", "run", now, now, now),
                    )
                client = TestClient(app)
                response = client.delete("/api/reports/run-delete")
                with db.connect() as conn:
                    deleted_at = conn.execute("select report_deleted_at from run_tasks where id = ?", ("run-delete",)).fetchone()[0]

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()["ok"])
        self.assertEqual(response.json()["run_id"], "run-delete")
        self.assertTrue(deleted_at)

    def test_deleted_report_detail_returns_not_found(self) -> None:
        try:
            from fastapi.testclient import TestClient
        except Exception as exc:  # pragma: no cover
            self.skipTest(f"TestClient unavailable: {exc}")

        with tempfile.TemporaryDirectory() as folder:
            db_path = Path(folder) / "test.sqlite3"
            with (
                patch("icm_platform.db.DB_PATH", db_path),
                patch("icm_platform.db.DATA_DIR", db_path.parent),
                patch("icm_platform.api.read_report", return_value="status: completed"),
                patch("icm_platform.api.parse_report", return_value={"case_id": "TC-004", "case_name": "Deleted detail", "status": "completed", "screenshots": []}),
            ):
                db.init_db()
                now = "2026-07-08T10:00:00Z"
                with db.connect() as conn:
                    conn.execute(
                        "insert into run_tasks(id, mode, case_id, status, command, created_at, started_at, finished_at, report_deleted_at) values (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                        ("run-hidden", "agent-explore", "TC-004", "completed", "run", now, now, now, now),
                    )
                client = TestClient(app)
                response = client.get("/api/reports/run-hidden")

        self.assertEqual(response.status_code, 404)
        self.assertEqual(response.json()["detail"], "report not found")


if __name__ == "__main__":
    unittest.main()
